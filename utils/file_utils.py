import re
from io import BytesIO
from mimetypes import guess_type
from pathlib import Path

import openpyxl
import xlrd
from docx import Document
from pypdf import PdfReader

from utils.oss_utils import upload_image_bytes

MAX_FILE_SIZE_BYTES = 10 * 1024 * 1024
MAX_TEXT_CHARS_PER_FILE = 90000
MAX_TEXT_CHARS_PER_PREVIEW = 1400
MAX_EXCEL_ROWS = 30
MAX_EXCEL_COLS = 12
MAX_CHUNK_CHARS = 1400
CHUNK_OVERLAP_CHARS = 180
MAX_PARAGRAPH_PREVIEW = 90

IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".webp", ".gif"}
TEXT_EXTENSIONS = {
    ".pdf",
    ".txt",
    ".md",
    ".csv",
    ".docx",
    ".doc",
    ".xlsx",
    ".xls",
}
SUPPORTED_EXTENSIONS = TEXT_EXTENSIONS | IMAGE_EXTENSIONS


class UnsupportedFileTypeError(Exception):
    pass


def _normalize_text(text: str) -> str:
    return "\n".join(line.rstrip() for line in text.replace("\r\n", "\n").splitlines()).strip()


def _truncate_text(text: str, suffix: str = "\n\n[内容已截断]") -> str:
    if len(text) <= MAX_TEXT_CHARS_PER_FILE:
        return text
    return text[:MAX_TEXT_CHARS_PER_FILE].rstrip() + suffix


def _read_txt(data: bytes) -> str:
    for encoding in ("utf-8", "utf-8-sig", "gb18030", "gbk"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def _read_pdf(data: bytes) -> str:
    reader = PdfReader(BytesIO(data))
    pages = []
    for index, page in enumerate(reader.pages, start=1):
        text = (page.extract_text() or "").strip()
        if text:
            pages.append(f"[第 {index} 页]\n{text}")
        else:
            pages.append(f"[第 {index} 页]\n[未提取到可读文本，可能是扫描件或图片页]")
    return "\n\n".join(pages)


def _read_docx(data: bytes) -> str:
    document = Document(BytesIO(data))
    blocks = []

    paragraph_blocks = []
    for index, paragraph in enumerate(document.paragraphs, start=1):
        text = paragraph.text.strip()
        if not text:
            continue
        preview = text[:MAX_PARAGRAPH_PREVIEW]
        preview_text = preview if preview == text else preview + "..."
        paragraph_blocks.append(f"[段落 {index}] {preview_text}\n{text}")
    if paragraph_blocks:
        blocks.append("\n\n".join(paragraph_blocks))

    for table_index, table in enumerate(document.tables, start=1):
        rows = []
        for row in table.rows[:MAX_EXCEL_ROWS]:
            cells = [cell.text.strip() for cell in row.cells[:MAX_EXCEL_COLS]]
            if any(cells):
                rows.append(" | ".join(cells))
        if rows:
            blocks.append(f"[表格 {table_index}]\n" + "\n".join(rows))

    return "\n\n".join(blocks)


def _read_xlsx(data: bytes) -> str:
    workbook = openpyxl.load_workbook(BytesIO(data), read_only=True, data_only=True)
    sheet_blocks = []
    for sheet in workbook.worksheets:
        rows = []
        header = None
        for row in sheet.iter_rows(max_row=MAX_EXCEL_ROWS, max_col=MAX_EXCEL_COLS, values_only=True):
            values = ["" if cell is None else str(cell).strip() for cell in row]
            if any(values):
                if header is None:
                    header = values
                rows.append("\t".join(values))
        if rows:
            block = [f"[Sheet: {sheet.title}]"]
            if header:
                block.append("列头: " + " | ".join(header))
            block.append("\n".join(rows))
            sheet_blocks.append("\n".join(block))
    return "\n\n".join(sheet_blocks)


def _read_xls(data: bytes) -> str:
    workbook = xlrd.open_workbook(file_contents=data)
    sheet_blocks = []
    for sheet in workbook.sheets():
        rows = []
        header = None
        for row_index in range(min(sheet.nrows, MAX_EXCEL_ROWS)):
            values = [
                str(sheet.cell_value(row_index, col_index)).strip()
                for col_index in range(min(sheet.ncols, MAX_EXCEL_COLS))
            ]
            if any(values):
                if header is None:
                    header = values
                rows.append("\t".join(values))
        if rows:
            block = [f"[Sheet: {sheet.name}]"]
            if header:
                block.append("列头: " + " | ".join(header))
            block.append("\n".join(rows))
            sheet_blocks.append("\n".join(block))
    return "\n\n".join(sheet_blocks)


def _guess_mime_type(filename: str) -> str:
    mime_type, _encoding = guess_type(filename)
    return mime_type or "application/octet-stream"


def _is_section_header(line: str) -> bool:
    stripped = line.strip()
    return bool(
        re.match(
            r"^\[(第\s*\d+\s*页|Sheet: .+|表格\s*\d+|段落\s*\d+)\]",
            stripped,
        )
    )


def _split_sections(text: str) -> list[tuple[str, str]]:
    sections = []
    current_label = "正文"
    current_lines = []
    for raw_line in text.splitlines():
        if _is_section_header(raw_line):
            body = "\n".join(current_lines).strip()
            if body:
                sections.append((current_label, body))
            current_label = raw_line.strip()
            current_lines = []
            continue
        current_lines.append(raw_line)

    body = "\n".join(current_lines).strip()
    if body:
        sections.append((current_label, body))
    return sections or [("正文", text)]


def _chunk_section(label: str, text: str) -> list[dict]:
    chunks = []
    cleaned = text.strip()
    if not cleaned:
        return chunks

    start = 0
    chunk_index = 1
    while start < len(cleaned):
        end = min(start + MAX_CHUNK_CHARS, len(cleaned))
        piece = cleaned[start:end].strip()
        if piece:
            chunks.append({"label": f"{label}#{chunk_index}", "text": piece})
            chunk_index += 1
        if end >= len(cleaned):
            break
        start = max(end - CHUNK_OVERLAP_CHARS, start + 1)
    return chunks


def build_text_chunks(text: str) -> list[dict]:
    normalized = _normalize_text(text)
    if not normalized:
        return []
    chunks = []
    for label, section_text in _split_sections(normalized):
        chunks.extend(_chunk_section(label, section_text))
    return chunks


def parse_file_bytes(filename: str, data: bytes, content_type: str | None = None) -> dict:
    suffix = Path(filename).suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        supported = ", ".join(sorted(SUPPORTED_EXTENSIONS))
        raise UnsupportedFileTypeError(f"{filename}，当前支持: {supported}")

    if len(data) > MAX_FILE_SIZE_BYTES:
        raise ValueError(f"{filename} 超过 10MB 大小限制")

    if suffix in IMAGE_EXTENSIONS:
        uploaded = upload_image_bytes(data, filename, content_type)
        return {
            "name": filename,
            "extension": suffix,
            "size_bytes": len(data),
            "note": "图像文件会直接发送给多模态模型，可用于 OCR 和图像理解。",
            "content": "",
            "preview": "",
            "chunks": [],
            "chunk_count": 0,
            "modality": "image",
            "mime_type": uploaded["content_type"],
            "data_url": uploaded["url"],
            "image_url": uploaded["url"],
            "storage": uploaded["storage"],
            "object_key": uploaded["object_key"],
        }

    note = ""
    if suffix in {".txt", ".md", ".csv"}:
        text = _read_txt(data)
    elif suffix == ".pdf":
        text = _read_pdf(data)
    elif suffix == ".docx":
        text = _read_docx(data)
    elif suffix == ".doc":
        text = "旧版 .doc 文件暂不支持可靠解析，请优先另存为 .docx 后重新上传。"
        note = "旧版 Word 格式仅返回兼容性提示"
    elif suffix == ".xlsx":
        text = _read_xlsx(data)
    elif suffix == ".xls":
        text = _read_xls(data)
    else:
        raise UnsupportedFileTypeError(f"{filename} 暂不支持")

    normalized = _normalize_text(text)
    if not normalized:
        normalized = "没有提取到可读文本，可能该文件主要是图片、扫描件或空内容。"

    truncated = _truncate_text(normalized)
    chunks = build_text_chunks(truncated)
    return {
        "name": filename,
        "extension": suffix,
        "size_bytes": len(data),
        "note": note,
        "content": truncated,
        "preview": truncated[:MAX_TEXT_CHARS_PER_PREVIEW].strip(),
        "chunks": chunks,
        "chunk_count": len(chunks),
        "modality": "text",
        "mime_type": _guess_mime_type(filename),
        "data_url": None,
        "image_url": None,
        "storage": "",
        "object_key": "",
    }


def parse_uploads(files) -> list[dict]:
    attachments = []
    for upload in files:
        if upload is None or not getattr(upload, "filename", None):
            continue
        data = upload.file.read()
        upload.file.seek(0)
        attachments.append(parse_file_bytes(upload.filename, data, getattr(upload, "content_type", None)))
    return attachments


def parse_local_file(path: str | Path) -> dict:
    target = Path(path)
    return parse_file_bytes(target.name, target.read_bytes(), None)
